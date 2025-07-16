import json
from docx import Document
import PyPDF2
import io
from docx2txt import process as docx2txt_process
from openai import OpenAI

class ResumeParser:
    def __init__(self):
        self.client = OpenAI(
            api_key='QcGCOyVichfHetzkUDeM:AUoiqAJtarlstnrJMcTI',
            base_url='https://spark-api-open.xf-yun.com/v1/'
        )
        
    def extract_text_from_file(self, file):
        """从文件中提取文本（支持PDF、DOCX、DOC）"""
        file_extension = file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(file)
        elif file_extension == 'docx':
            return self.extract_text_from_docx(file)
        elif file_extension == 'doc':
            return self.extract_text_from_doc(file)
        else:
            raise ValueError(f"不支持的文件格式: {file_extension}")
    
    def extract_text_from_pdf(self, file):
        """从PDF文件中提取文本"""
        try:
            # 确保文件指针在开始位置
            file.seek(0)
            # 读取PDF文件
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
            text = ""
            
            # 提取每一页的文本
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text += page_text + "\n"
            
            result = text.strip()
            return result if result else "PDF文件似乎为空或无法读取文本内容"
            
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")
        
    def extract_text_from_docx(self, file):
        """从DOCX文档中提取文本"""
        try:
            # 确保文件指针在开始位置
            file.seek(0)
            doc = Document(file)
            full_text = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():  # 只添加非空段落
                    full_text.append(para.text.strip())
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            result = '\n'.join(full_text)
            
            # 如果没有提取到文本，尝试其他方法
            if not result.strip():
                # 重置文件指针
                file.seek(0)
                # 尝试用docx2txt作为备用方案
                try:
                    # 将文件内容读入BytesIO对象
                    file_content = file.read()
                    file_like = io.BytesIO(file_content)
                    result = docx2txt_process(file_like)
                except Exception as docx2txt_error:
                    result = f"无法提取DOCX文件内容，主要方法和备用方法都失败了：{str(docx2txt_error)}"
            
            return result if result.strip() else "DOCX文件似乎为空或无法读取内容"
            
        except Exception as e:
            # 如果Document方法失败，尝试docx2txt
            try:
                file.seek(0)
                file_content = file.read()
                file_like = io.BytesIO(file_content)
                result = docx2txt_process(file_like)
                return result if result.strip() else "无法从DOCX文件提取文本内容"
            except Exception as e2:
                raise Exception(f"DOCX文件读取失败: {str(e)}，备用方法也失败: {str(e2)}")
    
    def extract_text_from_doc(self, file):
        """从DOC文档中提取文本"""
        try:
            # 确保文件指针在开始位置
            file.seek(0)
            # 对于老版本的DOC文件，使用docx2txt处理
            text = docx2txt_process(file)
            result = text.strip() if text else ""
            return result if result else "DOC文件似乎为空，请尝试转换为DOCX或PDF格式"
        except Exception as e:
            raise Exception(f"DOC文件读取失败: {str(e)}")
        
    # 保持向后兼容性
    def extract_text_from_word(self, file):
        """从Word文档中提取文本（保持向后兼容）"""
        return self.extract_text_from_docx(file)
    
    def analyze_resume(self, text):
        """分析简历内容 - 流式输出"""
        evaluation_criteria = """
        请根据以下标准分析这份简历：
        1. 技能匹配（35分）- 工作是否需要我的关键技术？
        2. 经验匹配（25分）- 我过去的角色与工作职责是否一致？
        3. 行业匹配度（15分）- 该工作是否与我曾工作过的行业相符？
        4. 公司与文化契合度（10分）- 是否符合我的工作风格（远程、敏捷、初创公司等）？
        5. 成长与兴趣（5分）- 这份工作是否让我兴奋或与我的职业目标一致？
        6. 可疑的简历信息。

        请给出详细的分析报告和各项评分。
        """

        try:
            response = self.client.chat.completions.create(
                model='generalv3.5',
                messages=[
                    {"role": "user", "content": f"简历内容：{text}\n\n{evaluation_criteria}"}
                ],
                stream=True  # 启用流式输出
            )
            return self._parse_streaming_response(response)
        except Exception as e:
            raise Exception(f"API调用失败: {str(e)}")
    
    def _parse_streaming_response(self, response):
        """解析流式响应"""
        full_response = ""
        try:
            for chunk in response:
                if chunk.choices[0].delta.content is not None:
                    content = chunk.choices[0].delta.content
                    full_response += content
        except Exception as e:
            raise Exception(f"流式响应解析失败: {str(e)}")
        return full_response
    
    def chat_with_ai(self, messages, new_message):
        """与AI进行对话 - 流式输出"""
        try:
            response = self.client.chat.completions.create(
                model='generalv3.5',
                messages=messages + [{"role": "user", "content": new_message}],
                stream=True  # 启用流式输出
            )
            return self._parse_streaming_response(response)
        except Exception as e:
            raise Exception(f"API调用失败: {str(e)}") 